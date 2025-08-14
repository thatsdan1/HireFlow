import logging
import PyPDF2
import docx
import io
import os
from typing import Dict, Any, Optional, List
from datetime import datetime

from app.services.ai_service import AIService

logger = logging.getLogger(__name__)

class ResumeService:
    def __init__(self):
        self.ai_service = AIService()
        self.allowed_types = ['.pdf', '.docx', '.doc', '.txt']
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    async def process_resume_upload(self, file_content: bytes, file_name: str, user_id: int) -> Dict[str, Any]:
        """Process uploaded resume file and extract structured data"""
        try:
            # Validate file
            self._validate_file(file_name, len(file_content))
            
            # Extract text based on file type
            text_content = self._extract_text(file_content, file_name)
            
            # Use AI to parse and structure the resume
            parsed_data = await self._ai_parse_resume(text_content)
            
            # Add metadata
            parsed_data['metadata'] = {
                'file_name': file_name,
                'file_size': len(file_content),
                'processed_at': datetime.now().isoformat(),
                'user_id': user_id
            }
            
            logger.info(f"Resume processed successfully for user {user_id}")
            
            return {
                'success': True,
                'text_content': text_content,
                'parsed_data': parsed_data,
                'file_name': file_name
            }
            
        except Exception as error:
            logger.error(f"Resume processing failed: {error}", exc_info=True)
            raise Exception(f"Resume processing failed: {str(error)}")
    
    def _validate_file(self, file_name: str, file_size: int) -> None:
        """Validate file type and size"""
        if file_size > self.max_file_size:
            raise Exception(f"File size {file_size} bytes exceeds maximum allowed size of {self.max_file_size} bytes")
        
        file_extension = os.path.splitext(file_name)[1].lower()
        if file_extension not in self.allowed_types:
            raise Exception(f"File type {file_extension} not supported. Allowed types: {', '.join(self.allowed_types)}")
    
    def _extract_text(self, file_content: bytes, file_name: str) -> str:
        """Extract text from different file types"""
        file_extension = os.path.splitext(file_name)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_content)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_text(file_content)
            elif file_extension == '.txt':
                return file_content.decode('utf-8')
            else:
                raise Exception(f"Unsupported file type: {file_extension}")
                
        except Exception as error:
            logger.error(f"Text extraction failed for {file_name}: {error}")
            raise Exception(f"Failed to extract text from {file_name}: {str(error)}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as page_error:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                    continue
            
            if not text.strip():
                raise Exception("No text could be extracted from PDF")
            
            return text.strip()
            
        except Exception as error:
            logger.error(f"PDF text extraction failed: {error}")
            raise Exception(f"Failed to extract text from PDF: {str(error)}")
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc = docx.Document(io.BytesIO(content))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if not text.strip():
                raise Exception("No text could be extracted from DOCX")
            
            return text.strip()
            
        except Exception as error:
            logger.error(f"DOCX text extraction failed: {error}")
            raise Exception(f"Failed to extract text from DOCX: {str(error)}")
    
    async def _ai_parse_resume(self, text_content: str) -> Dict[str, Any]:
        """Use AI to parse and structure resume text"""
        try:
            system_prompt = """You are an expert resume parser. Your task is to extract structured information from resume text and organize it into clear sections.

IMPORTANT RULES:
1. Extract ONLY information that is explicitly stated in the text
2. Do NOT invent or assume any details
3. Organize information into logical sections
4. Preserve the original wording and details
5. If information is missing for a section, leave it empty
6. Be precise and accurate with all extracted data"""

            user_prompt = f"""Please parse this resume text and extract structured information:

RESUME TEXT:
{text_content}

Please organize the information into these sections:
1. Personal Information (name, email, phone, location, linkedin, github)
2. Education (school, degree, gpa, graduation date, relevant coursework)
3. Experience (company, position, duration, location, description)
4. Skills (technical skills, soft skills, tools, languages)
5. Projects (name, description, technologies, outcomes)
6. Activities & Leadership (organizations, roles, achievements)
7. Awards & Honors (certifications, awards, recognition)

Return the information in a structured format that can be easily processed."""

            ai_response = self.ai_service._get_ai_response(system_prompt, user_prompt, {'resume_text': text_content})
            
            # Parse AI response into structured data
            parsed_data = self._parse_ai_response(ai_response)
            
            # Fallback to basic parsing if AI fails
            if not self._validate_parsed_data(parsed_data):
                logger.warning("AI parsing failed, using fallback parsing")
                parsed_data = self._fallback_parse(text_content)
            
            return parsed_data
            
        except Exception as error:
            logger.error(f"AI parsing failed: {error}")
            # Use fallback parsing
            return self._fallback_parse(text_content)
    
    def _parse_ai_response(self, ai_response: str) -> Dict[str, Any]:
        """Parse AI response into structured data"""
        try:
            # This is a simplified parser - in production, you'd want more sophisticated parsing
            parsed_data = {
                'personal': {},
                'education': [],
                'experience': [],
                'skills': [],
                'projects': [],
                'activities': [],
                'awards': []
            }
            
            # Basic parsing logic - look for section headers and extract content
            lines = ai_response.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect sections
                if 'personal' in line.lower() or 'contact' in line.lower():
                    current_section = 'personal'
                elif 'education' in line.lower():
                    current_section = 'education'
                elif 'experience' in line.lower() or 'work' in line.lower():
                    current_section = 'experience'
                elif 'skills' in line.lower():
                    current_section = 'skills'
                elif 'projects' in line.lower():
                    current_section = 'projects'
                elif 'activities' in line.lower() or 'leadership' in line.lower():
                    current_section = 'activities'
                elif 'awards' in line.lower() or 'honors' in line.lower():
                    current_section = 'awards'
                elif current_section and line:
                    # Add content to current section
                    self._add_to_section(parsed_data, current_section, line)
            
            return parsed_data
            
        except Exception as error:
            logger.error(f"Failed to parse AI response: {error}")
            return self._get_empty_structure()
    
    def _add_to_section(self, parsed_data: Dict[str, Any], section: str, content: str) -> None:
        """Add content to the appropriate section"""
        if section == 'personal':
            if '@' in content and '.' in content:
                parsed_data['personal']['email'] = content
            elif any(char.isdigit() for char in content) and any(char in '()-' for char in content):
                parsed_data['personal']['phone'] = content
            elif not parsed_data['personal'].get('name'):
                parsed_data['personal']['name'] = content
        
        elif section == 'education':
            if any(keyword in content.lower() for keyword in ['university', 'college', 'school']):
                parsed_data['education'].append({
                    'school': content,
                    'degree': '',
                    'gpa': '',
                    'graduation_date': '',
                    'relevant': ''
                })
        
        elif section == 'experience':
            if any(keyword in content.lower() for keyword in ['developer', 'engineer', 'analyst', 'manager']):
                parsed_data['experience'].append({
                    'position': content,
                    'company': '',
                    'duration': '',
                    'location': '',
                    'description': ''
                })
        
        elif section == 'skills':
            if ',' in content or '•' in content:
                skills = [skill.strip() for skill in content.replace('•', ',').split(',')]
                for skill in skills:
                    if skill and len(skill) > 2:
                        parsed_data['skills'].append({
                            'name': skill,
                            'level': 'Intermediate',
                            'category': 'Technical'
                        })
    
    def _validate_parsed_data(self, parsed_data: Dict[str, Any]) -> bool:
        """Validate that parsed data has meaningful content"""
        if not parsed_data:
            return False
        
        # Check if we have at least some personal info and one other section
        has_personal = bool(parsed_data.get('personal'))
        has_other_content = any([
            parsed_data.get('education'),
            parsed_data.get('experience'),
            parsed_data.get('skills')
        ])
        
        return has_personal and has_other_content
    
    def _fallback_parse(self, text_content: str) -> Dict[str, Any]:
        """Fallback parsing when AI fails"""
        try:
            lines = text_content.split('\n')
            parsed_data = self._get_empty_structure()
            
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect sections
                if any(keyword in line.lower() for keyword in ['education', 'academic']):
                    current_section = 'education'
                    continue
                elif any(keyword in line.lower() for keyword in ['experience', 'work', 'employment']):
                    current_section = 'experience'
                    continue
                elif any(keyword in line.lower() for keyword in ['skills', 'technical', 'competencies']):
                    current_section = 'skills'
                    continue
                elif any(keyword in line.lower() for keyword in ['projects', 'portfolio']):
                    current_section = 'projects'
                    continue
                elif any(keyword in line.lower() for keyword in ['name', 'email', 'phone', 'address']):
                    current_section = 'personal'
                    continue
                
                # Parse content based on current section
                if current_section == 'personal':
                    if '@' in line and '.' in line:
                        parsed_data['personal']['email'] = line
                    elif any(char.isdigit() for char in line) and any(char in '()-' for char in line):
                        parsed_data['personal']['phone'] = line
                    elif 'university' in line.lower() or 'college' in line.lower():
                        parsed_data['personal']['location'] = line
                    elif not parsed_data['personal'].get('name'):
                        parsed_data['personal']['name'] = line
                
                elif current_section == 'education':
                    if any(keyword in line.lower() for keyword in ['university', 'college', 'school']):
                        parsed_data['education'].append({
                            'school': line,
                            'degree': '',
                            'gpa': '',
                            'graduation_date': '',
                            'relevant': ''
                        })
                
                elif current_section == 'experience':
                    if any(keyword in line.lower() for keyword in ['developer', 'engineer', 'analyst', 'manager']):
                        parsed_data['experience'].append({
                            'position': line,
                            'company': '',
                            'duration': '',
                            'location': '',
                            'description': ''
                        })
                
                elif current_section == 'skills':
                    if ',' in line or '•' in line:
                        skills = [skill.strip() for skill in line.replace('•', ',').split(',')]
                        for skill in skills:
                            if skill and len(skill) > 2:
                                parsed_data['skills'].append({
                                    'name': skill,
                                    'level': 'Intermediate',
                                    'category': 'Technical'
                                })
            
            return parsed_data
            
        except Exception as error:
            logger.error(f"Fallback parsing failed: {error}")
            return self._get_empty_structure()
    
    def _get_empty_structure(self) -> Dict[str, Any]:
        """Return empty resume structure"""
        return {
            'personal': {},
            'education': [],
            'experience': [],
            'skills': [],
            'projects': [],
            'activities': [],
            'awards': []
        }
    
    async def enhance_resume_section(self, section_type: str, current_content: str, 
                                   target_role: str, company: str) -> Dict[str, Any]:
        """Enhance a specific resume section using AI"""
        try:
            system_prompt = """You are an expert resume writer. Your task is to enhance resume content to better match job requirements.

IMPORTANT RULES:
1. Use ONLY the information provided by the user
2. Do NOT fabricate or invent skills, experiences, or achievements
3. Improve language and structure while maintaining truthfulness
4. Add relevant keywords naturally from the job requirements
5. Focus on quantifiable achievements where possible
6. Use strong action verbs and professional language"""

            user_prompt = f"""Please enhance this {section_type} section for a {target_role} position at {company}:

CURRENT CONTENT:
{current_content}

Enhance this section to:
1. Better match the job requirements
2. Use stronger action verbs
3. Include relevant keywords naturally
4. Improve overall impact and readability
5. Maintain all factual information

Provide the enhanced content and a list of specific improvements made."""

            ai_response = self.ai_service._get_ai_response(system_prompt, user_prompt, {
                'section_type': section_type,
                'current_content': current_content,
                'target_role': target_role,
                'company': company
            })
            
            return {
                'enhanced_content': ai_response,
                'improvements': self._extract_improvements(ai_response),
                'section_type': section_type
            }
            
        except Exception as error:
            logger.error(f"Section enhancement failed: {error}")
            raise Exception(f"Failed to enhance {section_type} section: {str(error)}")
    
    def _extract_improvements(self, ai_response: str) -> List[str]:
        """Extract improvement suggestions from AI response"""
        improvements = []
        lines = ai_response.split('\n')
        
        for line in lines:
            line = line.strip()
            if any(keyword in line.lower() for keyword in ['improved', 'enhanced', 'added', 'changed', 'updated']):
                improvements.append(line)
        
        return improvements[:5]  # Return first 5 improvements 